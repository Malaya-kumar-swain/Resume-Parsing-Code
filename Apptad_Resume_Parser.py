#!/usr/bin/env python
# coding: utf-8

import os
import re
import io
import zipfile
import spacy
import pandas as pd
from docx import Document
from PIL import Image, UnidentifiedImageError
import pytesseract
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font, Alignment

# Load your custom-trained model
nlp_custom = spacy.load("ner_model")  # This loads the NER model you trained earlier

# Define a list of SAS technologies to search for
technologies = [
    "SAS Enterprise Guide", "SAS EG", "BASE SAS", "EG SAS", 
    "SAS DI Studio", "SAS Management Console", "SAS BI Dashboard", 
    "SAS VIYA", "SAS", "SAS AML"
]

# Define the list of trailing words to remove
trailing_words = [
    "tpm apptad", "10yrs servicenow", "sas support", "45 yrs", "10y 0m", "sap ds ecolab", 
    "sas", "sas apptad", "sas di", "sas developer", "sas di support", "it recruiter", "updated", 
    "mdm developer", "sas developer bangalore", "java mdm", "informatica idq mdm", "data quality", 
    "mdm admin", "sql developer", "sas di pune", "45 yrs", "6y 4m", "sas di mumbai", "9y 0m", "p360", 
    "it recruiter", "sql", "sas developer", "technical recruiter noida", "sas apptad", "sas di", 
    "python developer", "sas di developer", "sas support", "sas di developer mumbai", 
    "informatica mdm bangalore", "11y 0m", "python", "mdm updated", "etl testing", "mdm java", 
    "mdm testing", "data scientist", "etl test lead", "bdm gurgaon", "mdm bsa apptad", "sql", 
    "idq developer bangalore", "azure data engineer", "powerbi developer apptad", 
    "mdm production support", "bo mdm1", "sas developer pune", "5y 8m", "technical recruiter noida", 
    "ux designer", "old","di support", "di", "developer","apptad","developer bangalore","immigration specialist",
    "di pune", "di mumbai","inormatica mdm bangalore","mdm","testing", "inormatica mdm bangalore","developer pune",
    "immediate contract over servicenow","it support greater noida","bangalore","pune","technical  noida","mumbai"," lwd july  india",
    "azure architect immediate dec snow","informatica  days data eco","power bi  immediate joiner servicenow","cosmos db lwd dec snow"," tesing"
]

# Updated function to clean and remove trailing numbers and unwanted characters
def clean_name(name):
    name = re.sub(r'\(.*?\)', '', name)  # Remove content inside parentheses
    name = re.sub(r'[^\w\s]', '', name)  # Remove special characters
    name = re.sub(r'\d+$', '', name)  # Remove trailing numbers
    name = re.sub(r'\s+', ' ', name).strip()  # Remove extra spaces
    return name

# Function to remove titles from the beginning or end of the name
def remove_titles(name):
    titles = ['Sr', 'Mr', 'Ms', 'Dr', 'Mrs', 'Miss', 'Prof', 'Eng', 'Architect', 'Technical Architect', 'Senior', 'Lead', 'Manager', 'Director', 'Chief']
    name_parts = name.split()
    while name_parts and name_parts[0] in titles:
        name_parts = name_parts[1:]
    while name_parts and name_parts[-1] in titles:
        name_parts = name_parts[:-1]
    return ' '.join(name_parts)

# Function to remove trailing words from the name
def remove_trailing_words(name):
    for trailing_word in trailing_words:
        name = re.sub(rf'\b{re.escape(trailing_word)}\b', '', name, flags=re.IGNORECASE).strip()
    return name

# Updated function to extract name using filename and content match, with trailing number removal
def extract_name(file_name, text):
    invalid_starting_words = ['resume', 'naukri', 'cv', 'profile', 'job', 'application', 'candidate']
    base_name_parts = [part for part in re.split(r'[_\s\-!!]', os.path.splitext(file_name)[0])
                       if part.lower() not in invalid_starting_words and not part.isdigit()]
    base_name = ' '.join(base_name_parts).strip()

    excluded_patterns = [
        r'\b(?:mobile|phone|contact|tel|email|fax|resume|cv|job|profile|candidate)\b', 
        r'\d{10,}',  # Exclude long numbers which might be phone numbers
        r'[\w\.-]+@[\w\.-]+',  # Exclude email addresses
        r'\b\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}\b',  # General phone number pattern
        r'linkedin\.com',  # Exclude LinkedIn URLs
        r'\bname\b',  # Exclude lines starting with 'Name'
        r'\bcurriculum vitae\b',  # Exclude 'Curriculum Vitae'
        r'\bdate\b',  # Exclude lines starting with 'Date'
    ]
    excluded_regex = re.compile('|'.join(excluded_patterns), re.IGNORECASE)

    base_name_cleaned = clean_name(base_name.lower())
    if base_name_cleaned:
        cleaned_base_name = remove_trailing_words(remove_titles(base_name_cleaned))
        if len(cleaned_base_name.split()) <= 5:  # Allow up to 5 words
            return cleaned_base_name

    # Use the custom-trained SpaCy model to extract names
    doc = nlp_custom(text)
    name_candidates = [ent.text for ent in doc.ents if ent.label_ == "PERSON" and len(ent.text.split()) <= 5]
    
    # Further refine names and select the best candidate
    if name_candidates:
        refined_candidates = [remove_trailing_words(remove_titles(clean_name(name))) for name in name_candidates if remove_titles(name)]
        refined_candidates = [name.strip() for name in refined_candidates if name]
        if refined_candidates:
            return refined_candidates[0]
    
    # If no name is found in the document, return the filename-based name
    return cleaned_base_name if base_name else "Unknown"

# Function to search for technology keywords
def search_technology(text):
    detected_technologies = []
    for tech in technologies:
        if re.search(r'\b' + re.escape(tech) + r'\b', text, re.IGNORECASE):
            detected_technologies.append(tech)
    return ', '.join(detected_technologies) if detected_technologies else "N/A"

# Function to remove content within parentheses and special characters
def clean_client_name(name):
    name = re.sub(r'\(.*?\)', '', name)
    name = re.sub(r'[^\w\s]', '', name)
    name = re.sub(r',\s*$', '', name)
    name = re.sub(r'\bs\b$', '', name)
    return name.strip()

# Function to extract client names from tables
def extract_client_names_from_tables(doc):
    client_names = set()
    
    for table in doc.tables:
        rows = [row.cells for row in table.rows]
        
        for row in rows:
            if len(row) > 1:
                header = row[0].text.strip().lower()
                client_value = row[1].text.strip()
                
                if 'client' in header:
                    if len(client_value.split()) < 6 and client_value:
                        cleaned_name = clean_client_name(client_value)
                        client_names.add(cleaned_name)
    
    return client_names

# Function to extract client names from text
def extract_client_names_from_text(full_text):
    client_names = set()
    client_keywords = [
        "Client", "Client:", "Client Name:", "Client Name",
        "Organisation", "Organisation:", "Organisation Name", "Organisation Name:"
    ]
    
    for keyword in client_keywords:
        pattern = re.compile(rf'{re.escape(keyword)}\s*[:\-]?\s*([^\n,:.]+)', re.IGNORECASE)
        matches = pattern.findall(full_text)
        for match in matches:
            client_name = clean_client_name(match.strip())
            if len(client_name.split()) < 6 and client_name:
                client_names.add(client_name)
    
    return client_names

# Function to filter out unwanted terms and extract valid client names
def filter_client_names(client_names):
    exclusion_keywords = [
        "requirement", "concept", "understanding", "business", 
        "analysis", "profile", "summary", "document", "expectation", 
        "scope", "timeline", "specification", "Tools & Environment",
        "management", "vision", "standards", "agents"
    ]
    
    filtered_names = set()
    for name in client_names:
        name_lower = name.lower()
        if not any(exclusion_word in name_lower for exclusion_word in exclusion_keywords):
            filtered_names.add(name)
    
    return filtered_names

# Function to extract all client names from text and tables
def extract_client_names(doc, technology):
    client_names = set()
    
    client_names.update(extract_client_names_from_tables(doc))
    
    full_text = "\n".join([para.text for para in doc.paragraphs])
    client_names.update(extract_client_names_from_text(full_text))
    
    client_names = filter_client_names(client_names)
    
    if not client_names:
        if "SAS" in technology:
            return "There is no Client Information mentioned in the resume"
        return "N/A"
    
    return ', '.join(sorted(client_names))

# Function to extract text from images in .docx files and display skipped images
def extract_text_from_images(docx_path):
    extracted_text = ""
    
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        for file_info in docx_zip.infolist():
            if file_info.filename.startswith('word/media/'):
                try:
                    with docx_zip.open(file_info) as image_file:
                        image_data = image_file.read()
                        image = Image.open(io.BytesIO(image_data))
                        extracted_text += pytesseract.image_to_string(image)
                except (IOError, OSError, UnidentifiedImageError) as e:
                    print(f"Skipping image in {docx_path}, file {file_info.filename}: {e}")

    return extracted_text

# Function to process a single resume
def process_resume(file_path):
    doc = Document(file_path)
    
    full_text = "\n".join([para.text for para in doc.paragraphs])
    image_text = extract_text_from_images(file_path)
    full_text += "\n" + image_text
    
    name = extract_name(os.path.basename(file_path), full_text)
    
    technology = search_technology(full_text)
    
    client_names = extract_client_names(doc, technology) if technology != "N/A" else "N/A"

    return {
        "Name": name,
        "Technology": technology,
        "Client Names": client_names,
        "File Name": os.path.basename(file_path)
    }

# Main function to process all resumes in the input directory
def process_resumes(input_directory, output_file):
    resume_data = []

    for filename in os.listdir(input_directory):
        if filename.endswith(".docx"):
            file_path = os.path.join(input_directory, filename)
            resume_info = process_resume(file_path)
            resume_data.append(resume_info)

    df = pd.DataFrame(resume_data)
    
    df = df[["Name", "Technology", "Client Names", "File Name"]]

    wb = Workbook()
    ws = wb.active
    ws.title = "Processed_Resume"
    
    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    summary_ws = wb.create_sheet(title="Summary", index=0)
    
    total_resumes = len(resume_data)
    
    summary_text = (
        f"Thanks for processing the Profiles.\n"
        f"Total number of processed resume = {total_resumes}\n"
        f"This is Version 1.0\n\n"
        f"Created by - Malaya Kumar Swain\n"
        f"Supported by - Anupam Anand\n"
        f"               Arunesh Raj"
    )
    
    summary_ws.merge_cells('A1:E10')
    
    cell = summary_ws['A1']
    cell.value = summary_text
    cell.font = Font(bold=True, size=20)
    cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter

        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        
        ws.column_dimensions[column].width = max_length + 2

    ws.auto_filter.ref = ws.dimensions

    wb.save(output_file)

# Set input directory and output file path
input_directory = '/Users/malayakumarswain/Desktop/AppTad/Python_Project/Resume_Processing/Input/Resume/'  # Replace with your input directory path
output_file = '/Users/malayakumarswain/Desktop/AppTad/Python_Project/Resume_Processing/Output/output_resumes.xlsx' # Replace with your output directory path

# Process resumes and generate the Excel file
process_resumes(input_directory, output_file)